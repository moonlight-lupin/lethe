"""Manual restore (the "Restore" tab): re-identify a document that was
tokenised OUTSIDE Lethe, so there is no vault job. We supply the token→value
map ourselves and rebuild the file with the same engine the Re-identify tab uses.

Covers: the token-scan regex, token-type inference (which names get filed into
the dictionary), and a build_restorer + redact_document round-trip on txt/docx.
"""
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document

from lethe import build_restorer, extract_text, redact_document
# the manual-restore helpers live in the UI module (pure functions, no UI state)
from app import _RESTORE_TOKEN_RE, _TYPED_TOKEN_RE, _restore_defaults


def scan(text):
    """Distinct tokens in document order (mirrors the panel's scan)."""
    seen, out = set(), []
    for m in _RESTORE_TOKEN_RE.finditer(text):
        if m.group(0) not in seen:
            seen.add(m.group(0))
            out.append(m.group(0))
    return out


# ---- token scanning --------------------------------------------------------
sample = ("[PERSON_001] of [COUNTERPARTY_001] met [CLIENT A].\n"
          "See footnote [1] and [PROJECT_002]. Email [EMAIL_001].")
found = scan(sample)
assert found == ["[PERSON_001]", "[COUNTERPARTY_001]", "[CLIENT A]",
                 "[1]", "[PROJECT_002]", "[EMAIL_001]"], found
print("scan:", found)

# ---- per-row defaults (Type pre-selected, whether "Save" starts ticked) ----
# Every token is OFFERED for saving via an explicit dropdown + checkbox; these
# are only the starting points the user can override — nothing is hard-skipped.
custom = ["PROJECT", "FUND"]
assert _restore_defaults("[PERSON_001]", custom) == ("PERSON", True)
assert _restore_defaults("[COUNTERPARTY_001]", custom) == ("COUNTERPARTY", True)
assert _restore_defaults("[PROJECT_002]", custom) == ("PROJECT", True)     # custom type
assert _restore_defaults("[OTHER_001]", custom) == ("OTHER", True)
assert _restore_defaults("[CLIENT A]", custom) == ("COUNTERPARTY", True)   # free-form: still offered
assert _restore_defaults("[EMAIL_001]", custom) == ("OTHER", False)        # pattern: Save off by default
assert _restore_defaults("[PHONE_001]", custom) == ("OTHER", False)
assert _restore_defaults("[1]", custom) == ("OTHER", False)                # footnote: Save off
# the left "include" box is off only for bare numeric footnote markers
assert "[1]"[1:-1].strip().isdigit()
assert not "[PERSON_001]"[1:-1].strip().isdigit()
print("per-row defaults OK")

# ---- txt round-trip --------------------------------------------------------
mapping = {"[PERSON_001]": "John Smith", "[COUNTERPARTY_001]": "Acme Capital Partners",
           "[CLIENT A]": "Meridian Holdings"}
restore = build_restorer(mapping)
out, hits = restore(sample)
assert "John Smith" in out and "Acme Capital Partners" in out and "Meridian Holdings" in out
assert "[PERSON_001]" not in out and "[COUNTERPARTY_001]" not in out
# untouched tokens (no value supplied) are left exactly as-is
assert "[1]" in out and "[PROJECT_002]" in out and "[EMAIL_001]" in out
print("txt restore hits:", hits, "| output:", out.split(chr(10))[0])

# ---- docx round-trip (format preserved, names come back) -------------------
doc = Document()
doc.add_paragraph("Memo: [PERSON_001] leads the deal with [COUNTERPARTY_001].")
tbl = doc.add_table(rows=2, cols=1)
tbl.rows[0].cells[0].text = "Counterparty"
tbl.rows[1].cells[0].text = "[COUNTERPARTY_001]"
buf = io.BytesIO(); doc.save(buf); docx_bytes = buf.getvalue()

scanned = scan(extract_text(docx_bytes, "docx"))
assert "[PERSON_001]" in scanned and "[COUNTERPARTY_001]" in scanned, scanned
out_bytes, ext, dhits = redact_document(docx_bytes, "docx",
                                        build_restorer({"[PERSON_001]": "John Smith",
                                                        "[COUNTERPARTY_001]": "Acme Capital Partners"}))
out_text = extract_text(out_bytes, "docx")
assert ext == ".docx"
assert "John Smith" in out_text and "Acme Capital Partners" in out_text
assert "[PERSON_001]" not in out_text and "[COUNTERPARTY_001]" not in out_text
print("docx restore ext:", ext, "hits:", dhits, "| names back, tokens gone, table redacted")

print("RESTORE OK — out-of-band tokens scanned, typed, and rebuilt in-format")
