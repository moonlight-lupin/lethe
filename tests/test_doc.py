import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from openpyxl import Workbook

from lethe import (Entity, assign_tokens, build_replacer, detect,
                   extract_text, redact_document)

ents = [Entity("John Smith", "PERSON", ["Smith"]),
        Entity("Acme Capital Partners", "COUNTERPARTY", ["Acme"])]

# ---- DOCX (paragraphs + a table) -------------------------------------------
doc = Document()
doc.add_paragraph("Memo regarding John Smith and Acme Capital Partners.")
t = doc.add_table(rows=2, cols=2)
t.rows[0].cells[0].text = "Counterparty"
t.rows[0].cells[1].text = "Contact"
t.rows[1].cells[0].text = "Acme"
t.rows[1].cells[1].text = "john.smith@acme.com"
buf = io.BytesIO(); doc.save(buf); docx_bytes = buf.getvalue()

text = extract_text(docx_bytes, "docx")
items = assign_tokens(detect(text, ents))
repl, t2r = build_replacer(items)
out, ext, hits = redact_document(docx_bytes, "docx", repl)
out_text = extract_text(out, "docx")
print("DOCX ext:", ext, "hits:", hits)
print("  leaks John Smith:", "John Smith" in out_text, "| leaks Acme:", "Acme" in out_text,
      "| has token:", "[PERSON_001]" in out_text and "[COUNTERPARTY_001]" in out_text)

# ---- XLSX ------------------------------------------------------------------
wb = Workbook(); ws = wb.active
ws.append(["Party", "Email", "Amount"])
ws.append(["Acme Capital Partners", "john.smith@acme.com", 1000])
ws.append(["Note", "John Smith approved", "=C2*2"])
buf = io.BytesIO(); wb.save(buf); xlsx_bytes = buf.getvalue()

text = extract_text(xlsx_bytes, "xlsx")
items = assign_tokens(detect(text, ents))
repl, t2r = build_replacer(items)
out, ext, hits = redact_document(xlsx_bytes, "xlsx", repl)
out_text = extract_text(out, "xlsx")
print("XLSX ext:", ext, "hits:", hits)
print("  leaks John Smith:", "John Smith" in out_text, "| leaks Acme:", "Acme" in out_text,
      "| has token:", "[PERSON_001]" in out_text)

# confirm formula survived
from openpyxl import load_workbook
wb2 = load_workbook(io.BytesIO(out))
print("  formula preserved:", wb2.active["C3"].value == "=C2*2")
