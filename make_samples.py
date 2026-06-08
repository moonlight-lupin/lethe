"""Generate sample .docx / .pdf / .xlsx files for testing the De-identifier UI."""
import os

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font
from fpdf import FPDF

os.makedirs("samples", exist_ok=True)

MEMO = [
    "Dear Mr John Smith,",
    "",
    "Acme Capital Partners (Acme) confirms the secondary transaction with "
    "Meridian Holdings Pte Ltd. The mandate was led by Priya Raman and "
    "counter-signed by Wibowo Santoso of Garuda Ventures.",
    "",
    "Please remit to the DBS Bank account. Queries: john.smith@acme.com or "
    "+65 6789 1234. Reference account 1234-5678-9012.",
    "",
    "Kind regards,",
    "Jane Doe",
]

# ---- .docx (paragraphs + a small table) -------------------------------------
doc = Document()
doc.add_heading("Transaction Confirmation", level=1)
for line in MEMO:
    doc.add_paragraph(line)
doc.add_paragraph("")
t = doc.add_table(rows=1, cols=3)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Counterparty", "Contact", "Email"
for cp, contact, email in [
    ("Acme Capital Partners", "John Smith", "john.smith@acme.com"),
    ("Meridian Holdings Pte Ltd", "Priya Raman", "priya@meridian.sg"),
    ("Garuda Ventures", "Wibowo Santoso", "wibowo@garuda.id"),
]:
    row = t.add_row().cells
    row[0].text, row[1].text, row[2].text = cp, contact, email
doc.save("samples/sample-memo.docx")

# ---- .pdf (ASCII-safe so the core PDF font is happy) ------------------------
pdf = FPDF()
pdf.add_page()
pdf.set_font("Helvetica", "B", 14)
pdf.cell(0, 10, "Transaction Confirmation", new_x="LMARGIN", new_y="NEXT")
pdf.ln(2)
pdf.set_font("Helvetica", size=11)
for line in MEMO:
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.epw, 7, line if line else " ")
pdf.output("samples/sample-memo.pdf")

# ---- .xlsx (a counterparty register, with a formula) ------------------------
wb = Workbook()
ws = wb.active
ws.title = "Counterparties"
headers = ["Counterparty", "Contact", "Email", "Phone", "Amount (USD)"]
ws.append(headers)
for c in ws[1]:
    c.font = Font(bold=True)
data = [
    ("Acme Capital Partners", "John Smith", "john.smith@acme.com", "+65 6789 1234", 1000000),
    ("Meridian Holdings Pte Ltd", "Priya Raman", "priya@meridian.sg", "+65 6111 2222", 750000),
    ("Garuda Ventures", "Wibowo Santoso", "wibowo@garuda.id", "+62 21 555 7788", 500000),
]
for r in data:
    ws.append(list(r))
ws["E5"] = "=SUM(E2:E4)"   # a formula, to prove formulas survive redaction
ws["D5"] = "Total"
wb.save("samples/sample-counterparties.xlsx")

print("Created:")
for f in ("sample-memo.docx", "sample-memo.pdf", "sample-counterparties.xlsx"):
    print("  samples/" + f)
